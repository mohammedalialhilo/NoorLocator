from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = ROOT / "NoorLocator_Deployment_Guide.md"
OUTPUT_PATH = ROOT / "NoorLocator_Deployment_Guide.docx"

INLINE_TOKEN_PATTERN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
NUMBERED_ITEM_PATTERN = re.compile(r"^\d+\.\s+")
BULLET_ITEM_PATTERN = re.compile(r"^- ")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True

    if "CodeBlock" not in document.styles:
        code_style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = document.styles["CodeBlock"]

    code_style.base_style = document.styles["Normal"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9.5)

    document.core_properties.title = "NoorLocator Deployment Guide"
    document.core_properties.subject = "Beginner-friendly Azure deployment guide for NoorLocator"


def add_inline_runs(paragraph, text: str) -> None:
    position = 0

    for match in INLINE_TOKEN_PATTERN.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token)

        position = match.end()

    if position < len(text):
        paragraph.add_run(text[position:])


def add_code_block(document: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return

    paragraph = document.add_paragraph(style="CodeBlock")
    for index, line in enumerate(code_lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        if index < len(code_lines) - 1:
            run.add_break()


def add_markdown_line(document: Document, line: str, is_first_heading: bool) -> bool:
    if line.startswith("# "):
        paragraph = document.add_paragraph(style="Title" if is_first_heading else "Heading 1")
        add_inline_runs(paragraph, line[2:].strip())
        return False

    if line.startswith("## "):
        paragraph = document.add_paragraph(style="Heading 1")
        add_inline_runs(paragraph, line[3:].strip())
        return False

    if line.startswith("### "):
        paragraph = document.add_paragraph(style="Heading 2")
        add_inline_runs(paragraph, line[4:].strip())
        return False

    if line.startswith("#### "):
        paragraph = document.add_paragraph(style="Heading 3")
        add_inline_runs(paragraph, line[5:].strip())
        return False

    if NUMBERED_ITEM_PATTERN.match(line):
        paragraph = document.add_paragraph(style="List Number")
        add_inline_runs(paragraph, NUMBERED_ITEM_PATTERN.sub("", line, count=1))
        return False

    if BULLET_ITEM_PATTERN.match(line):
        paragraph = document.add_paragraph(style="List Bullet")
        add_inline_runs(paragraph, line[2:].strip())
        return False

    if line.strip():
        paragraph = document.add_paragraph(style="Normal")
        add_inline_runs(paragraph, line.strip())

    return False


def convert_markdown_to_docx(source_path: Path, output_path: Path) -> None:
    if not source_path.exists():
        raise FileNotFoundError(f"Source markdown file was not found: {source_path}")

    document = Document()
    configure_document(document)

    is_first_heading = True
    in_code_block = False
    code_lines: list[str] = []

    for raw_line in source_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code_block:
                add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if line.startswith("# "):
            add_markdown_line(document, line, is_first_heading)
            is_first_heading = False
            continue

        add_markdown_line(document, line, is_first_heading)

    if code_lines:
        add_code_block(document, code_lines)

    document.save(output_path)


if __name__ == "__main__":
    convert_markdown_to_docx(SOURCE_PATH, OUTPUT_PATH)
    print(f"Generated {OUTPUT_PATH}")
